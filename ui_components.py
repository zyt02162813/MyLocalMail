# ui_components.py
import os
import shutil
import smtplib
import re # 🔥 核心回归：用于正则提取会议链接
import sqlite3
from datetime import datetime
from email.utils import formataddr, parsedate_to_datetime
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from email.mime.application import MIMEApplication

from PyQt6.QtWidgets import (QWidget, QHBoxLayout, QVBoxLayout, QLabel, 
                             QPushButton, QFrame, QGridLayout, QSizePolicy, 
                             QCalendarWidget, QAbstractItemView, QDialog, 
                             QComboBox, QLineEdit, QTextEdit, QToolBar, QMessageBox, 
                             QFileDialog, QGraphicsDropShadowEffect, QGraphicsOpacityEffect, 
                             QToolButton, QCheckBox, QMenu)

from PyQt6.QtCore import Qt, QSize, QRect, QUrl, QDate, QPropertyAnimation, QEasingCurve, QPoint, pyqtSignal
from PyQt6.QtGui import QColor, QPainter, QFont, QFontMetrics, QAction, QTextCharFormat, QDesktopServices

import config

STYLESHEET = """
QMainWindow { background-color: #F5F7FA; }

/* === 工具栏 === */
QWidget#Toolbar { background-color: #FFFFFF; border-bottom: 1px solid #E5E5E5; padding: 12px 20px; }
QPushButton { border-radius: 6px; padding: 6px 16px; font-size: 13px; border: none; font-weight: 600; font-family: ".AppleSystemUIFont", "Segoe UI", sans-serif; }
QPushButton#ComposeBtn { background-color: #007AFF; color: white; }
QPushButton#ComposeBtn:hover { background-color: #0062cc; }
QPushButton#SyncBtn { background-color: #FFFFFF; color: #333; border: 1px solid #E0E0E0; }
QPushButton#SyncBtn:hover { background-color: #F9F9F9; }

/* === Tab === */
QTabWidget::pane { border: none; background: transparent; }
QTabBar { alignment: center; }
QTabBar::tab { background: transparent; color: #666; padding: 6px 20px; font-size: 14px; font-weight: 600; margin: 4px; border-radius: 6px; min-width: 80px; }
QTabBar::tab:selected { background-color: #E5E5EA; color: #000; }
QTabBar::tab:hover { color: #333; background-color: #F0F0F5; }

/* === 左侧 === */
QWidget#LeftPanel { background-color: #F7F8FA; border-right: 1px solid #EAEAEA; }
QWidget#SearchContainer { background-color: transparent; padding: 12px 10px 8px 10px; }
QFrame#SearchBoxFrame { background-color: #FFFFFF; border: 1px solid #E0E0E0; border-radius: 8px; min-height: 32px; max-height: 32px; }
QFrame#SearchBoxFrame:hover { border: 1px solid #C0C0C0; }
QLineEdit#TransparentSearchInput { background: transparent; border: none; font-size: 12px; color: #333; padding-left: 4px; }
QToolButton#InnerFilterBtn { background: transparent; border: none; border-radius: 6px; color: #888; padding: 1px 4px; }
QToolButton#InnerFilterBtn:hover { background-color: rgba(0,0,0,0.05); color: #333; }
QToolButton#InnerFilterBtn:checked { color: #007AFF; background-color: #EBF5FF; }

/* === 列表 === */
QTreeWidget { background-color: transparent; border: none; font-size: 13px; color: #444; outline: none; padding-top: 5px; }
QTreeWidget::item { height: 34px; border: none; border-radius: 6px; padding-left: 5px; margin: 1px 8px; }
QTreeWidget::item:selected { background-color: #E3E5E8; color: #000; font-weight: 600; }
QListWidget { background-color: white; border: none; border-right: 1px solid #EAEAEA; outline: none; }
QListWidget::item { border-bottom: 1px solid #F5F5F5; margin: 0px; padding: 0px; }
QListWidget::item:selected { background-color: #007AFF; }

/* === 🔥🔥🔥 日历组件专属样式 === */
QFrame#CalendarContainer { background-color: white; border-radius: 16px; border: 1px solid #EAEAEA; }
QFrame#EventCardFrame { 
    background-color: #FFFFFF; 
    border: 1px solid #E5E5E5; 
    border-radius: 12px; 
    margin-bottom: 12px; 
    padding: 0px; 
}
QFrame#EventCardFrame:hover { border-color: #007AFF; background-color: #FAFCFF; }
QLabel#EvtTime { color: #007AFF; font-weight: 700; font-size: 13px; font-family: monospace; }
QLabel#EvtTitle { color: #111; font-weight: 700; font-size: 16px; margin-bottom: 4px; }
QLabel#EvtLoc { color: #666; font-size: 12px; }
QLabel#EvtPeople { color: #888; font-size: 12px; margin-top: 4px; }

/* 🔥 会议链接按钮 */
QPushButton#MeetingBtn { 
    border-radius: 6px; 
    padding: 6px 12px; 
    font-size: 12px; 
    font-weight: 600; 
    color: white; 
    border: none; 
    text-align: left; 
    margin-top: 5px;
}
QPushButton#MeetingBtn:hover { opacity: 0.9; }

/* 🔥 纪要编辑器 */
QTextEdit#MinutesEditor { 
    background-color: #FAFAFA; 
    border: 1px solid #EEEEEE; 
    border-radius: 8px; 
    color: #333; 
    font-size: 13px; 
    line-height: 1.5;
    padding: 10px;
    margin-top: 5px;
}
QTextEdit#MinutesEditor:focus { background-color: #FFFFFF; border: 1px solid #007AFF; }

/* 🔥 纪要操作按钮 */
QPushButton#MinutesActionBtn { 
    background: transparent; 
    color: #888; 
    border: none; 
    font-size: 12px; 
    padding: 2px 6px; 
}
QPushButton#MinutesActionBtn:hover { color: #007AFF; background-color: #EBF5FF; border-radius: 4px; }
QLabel#MinutesStatus { color: #AAA; font-size: 11px; margin-right: 8px; }

/* === 邮件阅读器 === */
QWidget#MailReader { background-color: white; }
QFrame#ReaderHeader { background-color: white; border-bottom: 1px solid #F0F0F0; padding: 25px 30px 15px 30px; }
QLabel#DetailSubject { font-size: 20px; font-weight: 700; color: #111; margin-bottom: 8px; line-height: 1.4; }
QLabel#DetailTime { font-size: 12px; color: #999; margin-top: 4px;}
QLabel#DetailAvatar { background-color: #E1E4E8; color: #555; font-size: 16px; font-weight: bold; border-radius: 20px; min-width: 40px; min-height: 40px; max-width: 40px; max-height: 40px; qproperty-alignment: AlignCenter;}
QLabel#DetailSenderName { font-size: 15px; font-weight: 700; color: #222; }
QLabel#DetailSenderEmail { font-size: 13px; color: #666; margin-left: 5px; }

/* === 组件 === */
QFrame#AttachmentChip { background-color: #F5F7FA; border: 1px solid #E5E5E5; border-radius: 6px; min-height: 28px; }
QFrame#AttachmentChip:hover { background-color: #EBF5FF; border: 1px solid #007AFF; cursor: pointer; }
QLabel#AttName { font-size: 12px; color: #333; font-weight: 500; background: transparent; }
QLabel#AttSize { font-size: 11px; color: #888; margin-left: 5px; background: transparent; }
QFrame#PersonChipFrame { background-color: #E5E5EA; border-radius: 10px; padding: 0px; margin: 0px; }
QFrame#PersonChipFrame:hover { background-color: #D1D1D6; cursor: pointer; }
QLabel#ChipName { color: #111; font-size: 12px; font-weight: 500; padding: 2px 8px; background: transparent; }
QFrame#PersonCard { background-color: white; border: 1px solid #E5E5E5; border-radius: 12px; }
QLabel#CardName { font-size: 16px; font-weight: bold; color: #333; }
QLabel#CardEmail { font-size: 13px; color: #888; }
QPushButton#CardActionBtn { background-color: #F2F2F7; border: none; border-radius: 6px; color: #333; font-size: 12px; padding: 6px 12px; font-weight: 600; }
QPushButton#CardActionBtn:hover { background-color: #E5E5EA; color: #007AFF; }

/* 滚动条 */
QScrollBar:vertical { border: none; background: transparent; width: 10px; margin: 0px; }
QScrollBar::handle:vertical { background: #C1C1C1; min-height: 40px; border-radius: 5px; margin: 2px; }
QScrollBar::handle:vertical:hover { background: #8E8E93; }
QScrollBar::add-line:vertical, QScrollBar::sub-line:vertical { height: 0px; background: none; }
QScrollBar::add-page:vertical, QScrollBar::sub-page:vertical { background: none; }

/* 筛选 */
QFrame#FilterPopupFrame { background-color: #FFFFFF; border: 1px solid #E0E0E0; border-radius: 12px; }
QCheckBox { spacing: 10px; font-size: 13px; color: #333; padding: 6px 10px; border-radius: 6px; }
QCheckBox:hover { background-color: #F5F7FA; }
QCheckBox::indicator { width: 16px; height: 16px; border-radius: 4px; border: 1px solid #C0C0C0; background-color: white; }
QCheckBox::indicator:hover { border-color: #999; }
QCheckBox::indicator:checked { background-color: #007AFF; border: 1px solid #007AFF; image: none; }

/* 日历导航 */
QPushButton#NavBtn { background-color: transparent; border: none; border-radius: 15px; font-weight: bold; color: #666; width: 30px; height: 30px; }
QPushButton#NavBtn:hover { background-color: #f2f2f7; color: #333; }
QPushButton#YearMonthBtn { background-color: #F2F4F8; border: none; border-radius: 8px; font-size: 15px; font-weight: bold; color: #333; padding: 6px 14px; margin: 0 4px; }
QPushButton#YearMonthBtn:hover { background-color: #E6E8EC; }
QCalendarWidget { background-color: transparent; border: none; }
QCalendarWidget QWidget#qt_calendar_navigationbar { background-color: transparent; min-height: 0px; max-height: 0px; }
QCalendarWidget QTableView { alternate-background-color: white; background-color: white; selection-background-color: white; selection-color: black; outline: none; }
"""

# ================= 组件类 =================
class PersonChip(QFrame):
    click_signal = pyqtSignal(str, str, QPoint)
    def __init__(self, raw_str, parent=None):
        super().__init__(parent)
        self.setObjectName("PersonChipFrame")
        self.name = raw_str.split('<')[0].replace('"', '').strip()
        self.email = raw_str.split('<')[1].replace('>', '').strip() if '<' in raw_str else raw_str
        if not self.name: self.name = self.email.split('@')[0] if '@' in self.email else "Unknown"
        layout = QHBoxLayout(self); layout.setContentsMargins(0,0,0,0)
        lbl = QLabel(self.name); lbl.setObjectName("ChipName"); layout.addWidget(lbl)
        self.setCursor(Qt.CursorShape.PointingHandCursor)
    def mousePressEvent(self, e):
        if e.button() == Qt.MouseButton.LeftButton: self.click_signal.emit(self.name, self.email, e.globalPosition().toPoint())

class PersonPopup(QWidget):
    action_signal = pyqtSignal(str, str)
    def __init__(self, name, email, parent=None):
        super().__init__(parent)
        self.setWindowFlags(Qt.WindowType.Popup | Qt.WindowType.FramelessWindowHint)
        self.setAttribute(Qt.WidgetAttribute.WA_TranslucentBackground)
        self.email = email
        container = QFrame(self); container.setObjectName("PersonCard")
        shadow = QGraphicsDropShadowEffect(self); shadow.setBlurRadius(20); shadow.setColor(QColor(0,0,0,30)); shadow.setOffset(0,5); container.setGraphicsEffect(shadow)
        ml = QVBoxLayout(self); ml.setContentsMargins(10,10,10,10); ml.addWidget(container)
        l = QVBoxLayout(container); l.setContentsMargins(20,20,20,20); l.setSpacing(10)
        h = QHBoxLayout()
        av = QLabel(name[0].upper() if name else "?")
        av.setStyleSheet("background-color:#E1E4E8;color:#555;border-radius:20px;font-size:20px;font-weight:bold;min-width:40px;min-height:40px;qproperty-alignment:AlignCenter;")
        il = QVBoxLayout(); il.setSpacing(2); 
        n = QLabel(name); n.setObjectName("CardName"); e = QLabel(email); e.setObjectName("CardEmail")
        il.addWidget(n); il.addWidget(e)
        h.addWidget(av); h.addLayout(il); l.addLayout(h)
        line = QFrame(); line.setFrameShape(QFrame.Shape.HLine); line.setStyleSheet("border-top:1px solid #F0F0F0;"); l.addWidget(line)
        bl = QHBoxLayout(); 
        b1 = QPushButton("✉️ 写邮件"); b1.setObjectName("CardActionBtn"); b1.clicked.connect(lambda: self.emit_action("compose"))
        b2 = QPushButton("🔍 往来邮件"); b2.setObjectName("CardActionBtn"); b2.clicked.connect(lambda: self.emit_action("history"))
        bl.addWidget(b1); bl.addWidget(b2); l.addLayout(bl)
    def emit_action(self, t): self.action_signal.emit(t, self.email); self.close()
    def paintEvent(self, e): pass

class AttachmentChip(QFrame):
    def __init__(self, filename, filepath, size_bytes, parent=None):
        super().__init__(parent); self.setObjectName("AttachmentChip"); self.filepath = os.path.abspath(filepath); self.filename = filename
        self.setCursor(Qt.CursorShape.PointingHandCursor); self.setToolTip(f"{filename}\n双击打开 | 右键另存为")
        l = QHBoxLayout(self); l.setContentsMargins(8,4,8,4); l.setSpacing(6)
        icon = QLabel("📎"); icon.setStyleSheet("color:#007AFF;background:transparent;"); l.addWidget(icon)
        nl = QLabel(filename); nl.setObjectName("AttName"); l.addWidget(nl)
        sz = "0B"
        try:
            s = int(size_bytes)
            if s<1024: sz=f"{s}B"
            elif s<1048576: sz=f"{s/1024:.1f}KB"
            else: sz=f"{s/1048576:.1f}MB"
        except: pass
        sl = QLabel(sz); sl.setObjectName("AttSize"); l.addWidget(sl)
    def mouseDoubleClickEvent(self, e):
        if os.path.exists(self.filepath): QDesktopServices.openUrl(QUrl.fromLocalFile(self.filepath))
        else: QMessageBox.warning(self, "错误", "文件不存在")
    def mousePressEvent(self, e):
        if e.button() == Qt.MouseButton.RightButton:
            m = QMenu(self); a = QAction("📥 另存为", self); a.triggered.connect(self.save); m.addAction(a); m.exec(e.globalPosition().toPoint())
    def save(self):
        p, _ = QFileDialog.getSaveFileName(self, "另存", self.filename)
        if p and os.path.exists(self.filepath): shutil.copy2(self.filepath, p); QMessageBox.information(self,"成功","已保存")

class SearchFilterPopup(QWidget):
    filterChanged = pyqtSignal()
    def __init__(self, parent=None):
        super().__init__(parent); self.setWindowFlags(Qt.WindowType.Popup|Qt.WindowType.FramelessWindowHint); self.setAttribute(Qt.WidgetAttribute.WA_TranslucentBackground)
        self.c = QFrame(self); self.c.setObjectName("FilterPopupFrame")
        sh = QGraphicsDropShadowEffect(self); sh.setBlurRadius(25); sh.setColor(QColor(0,0,0,40)); sh.setOffset(0,6); self.c.setGraphicsEffect(sh)
        l = QVBoxLayout(self); l.setContentsMargins(15,15,15,15); l.addWidget(self.c)
        i = QVBoxLayout(self.c); i.setContentsMargins(15,15,15,15); i.setSpacing(10)
        lbl = QLabel("搜索范围"); lbl.setStyleSheet("color:#999;font-size:11px;font-weight:600;margin-bottom:5px;"); i.addWidget(lbl)
        self.cb1 = QCheckBox("标题"); self.cb1.setChecked(True); self.cb1.stateChanged.connect(self.filterChanged.emit)
        self.cb2 = QCheckBox("正文"); self.cb2.setChecked(True); self.cb2.stateChanged.connect(self.filterChanged.emit)
        self.cb3 = QCheckBox("发件人"); self.cb3.setChecked(True); self.cb3.stateChanged.connect(self.filterChanged.emit)
        i.addWidget(self.cb1); i.addWidget(self.cb2); i.addWidget(self.cb3)
    def paintEvent(self, e): pass

class ToastOverlay(QLabel):
    def __init__(self, parent, text):
        super().__init__(parent); self.setText(text); self.setAlignment(Qt.AlignmentFlag.AlignCenter)
        self.setStyleSheet("background-color:rgba(30,30,30,230);color:white;border-radius:10px;padding:12px 24px;font-size:14px;font-weight:600;")
        self.adjustSize(); self.raise_(); self.show()
        p = parent.rect(); self.move(int(p.width()/2-self.width()/2), int(p.height()-120))
        self.op = QGraphicsOpacityEffect(self); self.setGraphicsEffect(self.op)
        self.an = QPropertyAnimation(self.op, b"opacity"); self.an.setDuration(2500); self.an.setStartValue(1.0); self.an.setEndValue(0.0); self.an.setEasingCurve(QEasingCurve.Type.InExpo)
        self.an.finished.connect(self.deleteLater); self.an.start()

def format_email_date(s):
    if not s: return ""
    try:
        dt = datetime.fromisoformat(s) if '-' in s and ':' in s else parsedate_to_datetime(s)
        if dt.tzinfo: dt = dt.astimezone()
        return dt.strftime("%H:%M") if dt.date() == datetime.now().date() else dt.strftime("%Y/%m/%d")
    except: return str(s)[:10]

class MailListCard(QFrame):
    def __init__(self, subject, sender, date_str, preview, parent=None):
        super().__init__(parent); self.setObjectName("MailCard")
        self.full_subject = subject if subject else "(无主题)"; self.full_preview = preview[:100].replace('\n',' ').strip() if preview else "无预览"
        l = QVBoxLayout(self); l.setContentsMargins(15,12,15,12); l.setSpacing(6)
        t = QHBoxLayout(); n = sender.split('<')[0].replace('"','').strip(); 
        self.ls = QLabel(n if n else sender); self.ls.setObjectName("MailSender"); self.ls.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Preferred)
        self.ld = QLabel(format_email_date(date_str)); self.ld.setObjectName("MailDate"); self.ld.setAlignment(Qt.AlignmentFlag.AlignRight)
        t.addWidget(self.ls); t.addWidget(self.ld); l.addLayout(t)
        self.lsub = QLabel(self.full_subject); self.lsub.setObjectName("MailSubject"); l.addWidget(self.lsub)
        self.lpre = QLabel(self.full_preview); self.lpre.setObjectName("MailPreview"); l.addWidget(self.lpre)
    def resizeEvent(self, e):
        super().resizeEvent(e); w = self.width()-30
        self.lsub.setText(QFontMetrics(self.lsub.font()).elidedText(self.full_subject, Qt.TextElideMode.ElideRight, w))
        self.lpre.setText(QFontMetrics(self.lpre.font()).elidedText(self.full_preview, Qt.TextElideMode.ElideRight, w))
    def set_selected(self, s):
        c1, c2, c3 = ("white","#EEE","#DDD") if s else ("#222","#444","#999")
        self.ls.setStyleSheet(f"font-weight:700;font-size:14px;color:{c1};")
        self.lsub.setStyleSheet(f"font-size:13px;color:{c2};margin-top:2px;font-weight:500;")
        self.ld.setStyleSheet(f"font-size:12px;color:{c3};"); self.lpre.setStyleSheet(f"font-size:12px;color:{c3};margin-top:4px;")

# 🔥🔥🔥 旗舰版 EventCard：集成一键入会、参会人注入、模版管理
class EventCard(QFrame):
    def __init__(self, uid, start_time, end_time, summary, location, description, minutes_content, sender, recipient, parent=None):
        super().__init__(parent); self.uid = uid; self.summary = summary; self.start_time = start_time; self.recipient = recipient
        self.setObjectName("EventCardFrame")
        
        layout = QVBoxLayout(self)
        layout.setContentsMargins(20, 20, 20, 20); layout.setSpacing(12)
        
        # 1. 顶部元数据
        s_t = start_time.split(' ')[1][:5]; e_t = end_time.split(' ')[1][:5]
        h_time = QHBoxLayout()
        h_time.addWidget(QLabel("🕒")); lbl_t = QLabel(f"{s_t} - {e_t}"); lbl_t.setObjectName("EvtTime"); h_time.addWidget(lbl_t); h_time.addStretch()
        layout.addLayout(h_time)
        
        lbl_title = QLabel(summary); lbl_title.setObjectName("EvtTitle"); lbl_title.setWordWrap(True)
        layout.addWidget(lbl_title)
        
        h_loc = QHBoxLayout()
        h_loc.addWidget(QLabel("📍")); lbl_l = QLabel(location if location else "线上/无地点"); lbl_l.setObjectName("EvtLoc"); h_loc.addWidget(lbl_l); h_loc.addStretch()
        layout.addLayout(h_loc)
        
        # 解析参会人 (用于显示和注入模版)
        self.attendees_list = []
        if recipient:
            for r in recipient.split(','):
                n = r.split('<')[0].replace('"','').strip()
                if n: self.attendees_list.append(n)
        
        h_ppl = QHBoxLayout()
        h_ppl.addWidget(QLabel("👥"))
        ppl_str = ", ".join(self.attendees_list[:3]) + ("..." if len(self.attendees_list)>3 else "") if self.attendees_list else "我"
        lbl_p = QLabel(f"参会: {ppl_str}"); lbl_p.setObjectName("EvtPeople"); h_ppl.addWidget(lbl_p); h_ppl.addStretch()
        layout.addLayout(h_ppl)
        
        # 2. 🔥 智能入会按钮
        link_url, link_type = self.detect_meeting_link(location, description)
        if link_url:
            btn_join = QPushButton()
            btn_join.setObjectName("MeetingBtn")
            btn_join.setCursor(Qt.CursorShape.PointingHandCursor)
            btn_join.clicked.connect(lambda: QDesktopServices.openUrl(QUrl(link_url)))
            
            if link_type == "Teams":
                btn_join.setText("🟣 Join Teams Meeting")
                btn_join.setStyleSheet("background-color: #464EB8; color: white;")
            elif link_type == "Tencent":
                btn_join.setText("🔵 加入腾讯会议")
                btn_join.setStyleSheet("background-color: #006EFF; color: white;")
            else:
                btn_join.setText("🔗 进入会议")
                btn_join.setStyleSheet("background-color: #2D8CFF; color: white;")
                
            layout.addWidget(btn_join)

        # 分割线
        line = QFrame(); line.setFrameShape(QFrame.Shape.HLine); line.setStyleSheet("border-top: 1px dashed #E0E0E0; margin-top: 10px; margin-bottom: 5px;")
        layout.addWidget(line)
        
        # 3. 纪要工具栏 (保存状态、重置、导出)
        h_tools = QHBoxLayout()
        h_tools.addWidget(QLabel("📝 会议纪要"))
        self.lbl_status = QLabel("☁️ 已同步"); self.lbl_status.setObjectName("MinutesStatus")
        h_tools.addWidget(self.lbl_status); h_tools.addStretch()
        
        btn_reset = QPushButton("↺ 重置"); btn_reset.setObjectName("MinutesActionBtn"); btn_reset.setCursor(Qt.CursorShape.PointingHandCursor)
        btn_reset.clicked.connect(self.reset_template)
        
        btn_export = QPushButton("📥 导出 Word"); btn_export.setObjectName("MinutesActionBtn"); btn_export.setCursor(Qt.CursorShape.PointingHandCursor)
        btn_export.clicked.connect(self.export_word)
        
        h_tools.addWidget(btn_reset); h_tools.addWidget(btn_export)
        layout.addLayout(h_tools)
        
        # 4. 编辑器
        self.ed = QTextEdit(); self.ed.setObjectName("MinutesEditor")
        self.ed.setMinimumHeight(150)
        self.ed.setPlaceholderText("在此记录会议纪要...")
        
        # 🔥 自动填充模版
        if minutes_content and len(minutes_content) > 10: # 简单判断是否有内容
            self.ed.setHtml(minutes_content)
        else:
            self.reset_template(confirm=False) # 初始加载默认模版
            
        self.ed.textChanged.connect(self.auto_save)
        layout.addWidget(self.ed)

    def detect_meeting_link(self, loc, desc):
        text = (str(loc) + " " + str(desc)).lower()
        # 简单的正则匹配
        url_match = re.search(r'(https?://[^\s]+)', text) # 粗略提取
        url = url_match.group(0) if url_match else None
        
        if "teams" in text: return url, "Teams"
        if "腾讯会议" in text or "voov" in text or "tencent" in text: return url, "Tencent"
        if "zoom" in text: return url, "Zoom"
        return None, None

    def reset_template(self, confirm=True):
        if confirm and QMessageBox.question(self, "重置", "确定要覆盖当前纪要吗？", QMessageBox.StandardButton.Yes|QMessageBox.StandardButton.No) != QMessageBox.StandardButton.Yes:
            return
            
        # 生成带参会人的 HTML 模版
        atts_html = ", ".join(self.attendees_list) if self.attendees_list else "暂无"
        html = f"""
        <p style='color:#666;'><b>👥 参会人：</b>{atts_html}</p>
        <hr>
        <p style='color:#333;'><b>✅ 结论 / 待办 (Action Items)：</b></p>
        <ul style='margin-top:0;'>
            <li>[ ] </li>
        </ul>
        <br>
        <p style='color:#333;'><b>💡 讨论要点：</b></p>
        <ul>
            <li> </li>
        </ul>
        """
        self.ed.setHtml(html)
        if confirm: self.auto_save() # 如果是手动重置，立即保存

    def auto_save(self):
        self.lbl_status.setText("💾 保存中...")
        try:
            conn = sqlite3.connect('local_mail.db'); c = conn.cursor()
            c.execute("UPDATE events SET minutes = ? WHERE uid = ?", (self.ed.toHtml(), self.uid))
            conn.commit(); conn.close()
            self.lbl_status.setText("☁️ 已同步")
        except: self.lbl_status.setText("❌ 失败")

    def export_word(self):
        try:
            from docx import Document # 需要安装 python-docx，如果没有会报错
            fname, _ = QFileDialog.getSaveFileName(self, "导出纪要", f"{self.summary}.docx")
            if not fname: return
            
            doc = Document()
            doc.add_heading(self.summary, 0)
            doc.add_paragraph(f"时间: {self.start_time}")
            doc.add_heading('会议纪要', level=1)
            doc.add_paragraph(self.ed.toPlainText()) # 简单导出纯文本
            doc.save(fname)
            QMessageBox.information(self, "成功", "导出完成！")
            QDesktopServices.openUrl(QUrl.fromLocalFile(fname))
        except ImportError:
            QMessageBox.warning(self, "错误", "请先安装: pip install python-docx")
        except Exception as e:
            QMessageBox.warning(self, "错误", str(e))

class MeetingCalendarWidget(QCalendarWidget):
    def __init__(self, parent=None):
        super().__init__(parent); self.data = {}; self.setMouseTracking(True); self.setNavigationBarVisible(False); self.setVerticalHeaderFormat(QCalendarWidget.VerticalHeaderFormat.NoVerticalHeader); self.setHeaderTextFormat(QTextCharFormat())
    def set_meeting_data(self, d): self.data = d; self.updateCell(QDate.currentDate())
    def mouseMoveEvent(self, e): super().mouseMoveEvent(e); self.update() 
    def paintCell(self, p, r, d):
        sel = d == self.selectedDate(); today = d == QDate.currentDate(); p.save(); p.setRenderHint(QPainter.RenderHint.Antialiasing); p.setPen(Qt.PenStyle.NoPen) 
        if sel: p.setBrush(QColor("#EBF5FF")); p.drawRoundedRect(r.adjusted(2,2,-2,-2),6,6); p.setPen(QColor("#0066FF")); p.setBrush(Qt.BrushStyle.NoBrush); p.drawRoundedRect(r.adjusted(2,2,-2,-2),6,6)
        else: p.setBrush(QColor("white")); p.drawRect(r)
        p.setPen(QColor("#0066FF") if today else (QColor("#333") if d.month() == self.monthShown() else QColor("#DDD")))
        if today: p.setBrush(QColor("#0066FF")); p.drawEllipse(r.center().x()-10, r.top()+4, 20, 20); p.setPen(QColor("white"))
        p.setFont(QFont("Arial", 10, QFont.Weight.Bold)); p.drawText(QRect(r.left(), r.top()+4, r.width(), 20), Qt.AlignmentFlag.AlignHCenter, str(d.day()))
        
        # 🔥🔥🔥 核心修复：文字绘制回归
        if self.data.get(d):
            bar_bg = QColor("#E3F2FD"); txt_col = QColor("#5F6368")
            bar_font = QFont("Arial", 9); p.setFont(bar_font); metrics = QFontMetrics(bar_font)
            start_y = r.top() + 28 
            for title in self.data[d][:3]: 
                bar_rect = QRect(r.left() + 3, start_y, r.width() - 6, 16)
                p.setBrush(bar_bg); p.setPen(Qt.PenStyle.NoPen); p.drawRoundedRect(bar_rect, 3, 3) 
                p.setPen(txt_col)
                text_rect = bar_rect.adjusted(4, 0, -2, 0)
                elided = metrics.elidedText(title, Qt.TextElideMode.ElideRight, text_rect.width())
                p.drawText(text_rect, Qt.AlignmentFlag.AlignLeft | Qt.AlignmentFlag.AlignVCenter, elided)
                start_y += 18 
        p.restore()

class ComposeWindow(QDialog):
    def __init__(self):
        super().__init__(); self.setWindowTitle("写邮件"); self.resize(900, 700); l = QVBoxLayout(self); self.att = [] 
        h1 = QHBoxLayout(); h1.addWidget(QLabel("发件人:")); self.c_from = QComboBox()
        for a in config.ACCOUNTS: self.c_from.addItem(f"{a['name']} <{a['email']}>", a)
        l.addLayout(h1); h1.addWidget(self.c_from)
        h2 = QHBoxLayout(); h2.addWidget(QLabel("收件人:")); self.i_to = QLineEdit(); l.addLayout(h2); h2.addWidget(self.i_to)
        h3 = QHBoxLayout(); h3.addWidget(QLabel("抄送 CC:")); self.i_sub = QLineEdit(); l.addLayout(h3); h3.addWidget(self.i_sub)
        h4 = QHBoxLayout(); h4.addWidget(QLabel("主   题:")); self.i_subject = QLineEdit(); l.addLayout(h4); h4.addWidget(self.i_subject)
        tb = QToolBar(); tb.setStyleSheet("border:none"); a = QAction("📎 附件", self); a.triggered.connect(self.add_att); tb.addAction(a); l.addWidget(tb)
        self.lbl_att = QLabel(""); self.lbl_att.setStyleSheet("color:blue"); l.addWidget(self.lbl_att)
        self.txt = QTextEdit(); self.txt.setStyleSheet("border:1px solid #ddd;padding:10px"); l.addWidget(self.txt)
        acc = self.c_from.currentData(); self.txt.setPlainText("\n" + acc.get('signature', ''))
        h5 = QHBoxLayout(); btn = QPushButton("🚀 发送"); btn.clicked.connect(self.send); h5.addStretch(); h5.addWidget(btn); l.addLayout(h5)
    def set_initial_data(self, to_addr="", subject=""):
        if to_addr: self.i_to.setText(to_addr)
        if subject: self.i_subject.setText(subject)
    def add_att(self):
        fs, _ = QFileDialog.getOpenFileNames(self); self.att.extend(fs); self.lbl_att.setText(f"附件: {len(self.att)} 个")
    def send(self):
        try:
            acc = self.c_from.currentData(); m = MIMEMultipart(); m['From'] = formataddr([acc['name'], acc['email']]); m['To'] = self.i_to.text(); m['Subject'] = self.i_subject.text()
            m.attach(MIMEText(self.txt.toPlainText(), 'plain', 'utf-8'))
            for f in self.att: 
                with open(f, 'rb') as x: part = MIMEApplication(x.read()); part.add_header('Content-Disposition', 'attachment', filename=os.path.basename(f)); m.attach(part)
            s = smtplib.SMTP_SSL(acc['smtp_server'], acc['smtp_port']); s.login(acc['email'], acc['password']); s.sendmail(acc['email'], self.i_to.text().split(','), m.as_string()); s.quit()
            QMessageBox.information(self, "成功", "已发送"); self.close()
        except Exception as e: QMessageBox.critical(self, "错误", str(e))